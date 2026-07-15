"""文档格式转换。

支持的转换路径:
    PDF  → DOCX, TXT, MD, HTML, RTF  (pdf2docx / PyPDF2)
    DOCX → PDF, TXT, HTML, MD, RTF   (python-docx / reportlab / weasyprint)
    DOC  → DOCX, PDF, TXT, HTML, MD, RTF  (LibreOffice chain)
    TXT  → PDF, DOCX, MD, HTML, RTF  (reportlab / python-docx)
    MD   → PDF, HTML, TXT, DOCX, RTF (markdown / weasyprint)
    HTML → PDF, TXT, DOCX, MD, RTF   (weasyprint / BeautifulSoup)
    EPUB → TXT, PDF, DOCX, MD, HTML, RTF (ebooklib)
    RTF  → TXT, PDF, DOCX, MD, HTML  (regex extraction / chain)
"""

from __future__ import annotations

import logging
import os
import re
import shutil
import tempfile
from typing import Callable

logger = logging.getLogger(__name__)


def convert_document(
    input_path: str,
    output_path: str,
    src_fmt: str,
    dst_fmt: str,
    progress_callback: Callable[[float], None] | None = None,
) -> None:
    """文档格式转换入口。"""
    key = (src_fmt, dst_fmt)

    handler = _HANDLERS.get(key)
    if handler is None:
        raise ValueError(f"不支持的文档转换: {src_fmt} → {dst_fmt}")

    handler(input_path, output_path)
    if progress_callback:
        progress_callback(1.0)


# ============================================================================
#  Text extraction helpers — each source format → plain text
# ============================================================================

def _read_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_pdf_text(input_path: str) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(input_path)
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx_text(input_path: str) -> str:
    from docx import Document
    doc = Document(input_path)
    return "\n".join(para.text for para in doc.paragraphs)


def _extract_html_text(input_path: str) -> str:
    from bs4 import BeautifulSoup
    html = _read_text_file(input_path)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text("\n")


def _extract_md_text(input_path: str) -> str:
    """Strip markdown syntax to get plain text."""
    text = _read_text_file(input_path)
    # Remove headers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove bold/italic
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}(.+?)_{1,3}", r"\1", text)
    # Remove links
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    # Remove images
    text = re.sub(r"!\[([^\]]*)\]\([^\)]+\)", r"\1", text)
    # Remove code blocks
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Remove blockquotes
    text = re.sub(r"^>\s+", "", text, flags=re.MULTILINE)
    # Remove list markers
    text = re.sub(r"^[\-\*\+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\d+\.\s+", "", text, flags=re.MULTILINE)
    # Remove horizontal rules
    text = re.sub(r"^[\-\*_]{3,}$", "", text, flags=re.MULTILINE)
    return text.strip()


def _extract_epub_text(input_path: str) -> str:
    from ebooklib import epub
    from bs4 import BeautifulSoup
    book = epub.read_epub(input_path)
    parts: list[str] = []
    for item in book.get_items_of_type(9):  # ITEM_DOCUMENT
        soup = BeautifulSoup(item.get_body_content(), "html.parser")
        parts.append(soup.get_text())
    return "\n\n".join(parts)


def _extract_rtf_text(input_path: str) -> str:
    content = _read_text_file(input_path)
    text = re.sub(r"[\\{].*?[\\}]", " ", content)
    text = re.sub(r"\\[a-z]+\d*\s?", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text).strip()
    return text


# ============================================================================
#  Text → target format helpers
# ============================================================================

def _save_as_txt(text: str, output_path: str) -> None:
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)


def _save_as_md(text: str, output_path: str) -> None:
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)


def _save_as_html(text: str, output_path: str, title: str = "Document") -> None:
    # Escape HTML
    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    paragraphs = "\n".join(f"<p>{line}</p>" for line in safe.split("\n") if line.strip())
    html = f"""<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="utf-8">
<title>{title}</title></head><body>
{paragraphs}
</body></html>"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)


def _save_as_rtf(text: str, output_path: str) -> None:
    # Basic RTF — escape special chars
    safe = text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
    safe = safe.replace("\n", "\\par\n")
    rtf = f"""{{\\rtf1\\ansi\\deff0
{{\\fonttbl{{\\f0 Times New Roman;}}}}
\\f0\\fs24
{safe}
}}"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(rtf)


def _save_text_as_pdf(text: str, output_path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story: list = []
    for line in text.split("\n"):
        if not line.strip():
            story.append(Spacer(1, 12))
        else:
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, styles["Normal"]))
    pdf_doc.build(story)


def _save_text_as_docx(text: str, output_path: str) -> None:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)


def _save_text_as_epub(text: str, output_path: str, title: str = "Converted Document") -> None:
    from ebooklib import epub
    book = epub.EpubBook()
    book.set_identifier("converted-doc")
    book.set_title(title)
    book.set_language("zh")
    # Split text into chapters by double newline
    chapters = [c.strip() for c in text.split("\n\n") if c.strip()]
    if not chapters:
        chapters = [text]
    spine = ["nav"]
    for i, ch_text in enumerate(chapters):
        chapter = epub.EpubHtml(title=f"Chapter {i+1}", file_name=f"chap_{i+1}.xhtml")
        paragraphs = "".join(f"<p>{line}</p>" for line in ch_text.split("\n") if line.strip())
        chapter.content = f"<html><body>{paragraphs}</body></html>"
        book.add_item(chapter)
        spine.append(chapter)
    book.toc = spine[1:]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine
    epub.write_epub(output_path, book, {})


# ============================================================================
#  Chain conversion helper — src → intermediate → dst
# ============================================================================

def _chain_via(intermediate_fmt: str):
    """Create a handler that converts src→intermediate→dst using a temp file."""
    def handler(input_path: str, output_path: str, _src: str = "", _dst: str = "") -> None:
        # Determine src and dst from file extensions
        src = _src or os.path.splitext(input_path)[1].lstrip(".").lower()
        dst = _dst or os.path.splitext(output_path)[1].lstrip(".").lower()
        tmp = tempfile.mktemp(suffix=f".{intermediate_fmt}")
        try:
            h1 = _HANDLERS.get((src, intermediate_fmt))
            h2 = _HANDLERS.get((intermediate_fmt, dst))
            if not h1 or not h2:
                raise ValueError(f"Cannot chain {src}→{intermediate_fmt}→{dst}")
            h1(input_path, tmp)
            h2(tmp, output_path)
        finally:
            if os.path.exists(tmp):
                os.unlink(tmp)
    return handler


# ============================================================================
#  Original conversion functions (kept as-is)
# ============================================================================

def _pdf_to_docx(input_path: str, output_path: str) -> None:
    from pdf2docx import Converter
    cv = Converter(input_path)
    cv.convert(output_path)
    cv.close()


def _pdf_to_txt(input_path: str, output_path: str) -> None:
    _save_as_txt(_extract_pdf_text(input_path), output_path)


def _pdf_to_image(input_path: str, output_path: str) -> None:
    from PyPDF2 import PdfReader
    from PIL import Image, ImageDraw, ImageFont
    reader = PdfReader(input_path)
    images: list[Image.Image] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        img = Image.new("RGB", (800, 1100), "white")
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("arial.ttf", 14)
        except Exception:
            font = ImageFont.load_default()
        y = 50
        for line in text.split("\n"):
            if y > 1050:
                break
            draw.text((50, y), line, fill="black", font=font)
            y += 20
        images.append(img)
    if len(images) == 1:
        images[0].save(output_path)
    else:
        images[0].save(output_path, save_all=True, append_images=images[1:])


def _docx_to_pdf(input_path: str, output_path: str) -> None:
    if _try_libreoffice(input_path, output_path, "pdf"):
        return
    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    doc = Document(input_path)
    pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story: list = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 12))
            continue
        try:
            story.append(Paragraph(text, styles["Normal"]))
        except Exception:
            safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, styles["Normal"]))
    pdf_doc.build(story)


def _docx_to_txt(input_path: str, output_path: str) -> None:
    _save_as_txt(_extract_docx_text(input_path), output_path)


def _docx_to_html(input_path: str, output_path: str) -> None:
    from docx import Document
    doc = Document(input_path)
    parts: list[str] = ["<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>"]
    for para in doc.paragraphs:
        text = para.text
        style = para.style.name if para.style else "Normal"
        tag = "h1" if "Heading 1" in style else ("h2" if "Heading 2" in style else "p")
        parts.append(f"<{tag}>{text}</{tag}>")
    parts.append("</body></html>")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(parts))


def _doc_to_docx(input_path: str, output_path: str) -> None:
    if _try_libreoffice(input_path, output_path, "docx"):
        return
    raise RuntimeError("DOC → DOCX 需要安装 LibreOffice。")


def _txt_to_pdf(input_path: str, output_path: str) -> None:
    _save_text_as_pdf(_read_text_file(input_path), output_path)


def _txt_to_docx(input_path: str, output_path: str) -> None:
    _save_text_as_docx(_read_text_file(input_path), output_path)


def _md_to_pdf(input_path: str, output_path: str) -> None:
    try:
        import markdown
        from weasyprint import HTML
        md_text = _read_text_file(input_path)
        html_body = markdown.markdown(md_text, extensions=["extra", "codehilite", "tables"])
        html_doc = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; line-height: 1.6; color: #333; }}
code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 4px; }}
pre {{ background: #f4f4f4; padding: 16px; border-radius: 8px; overflow-x: auto; }}</style></head><body>{html_body}</body></html>"""
        HTML(string=html_doc).write_pdf(output_path)
    except ImportError:
        raise RuntimeError("MD → PDF 需要安装 weasyprint。")


def _md_to_html(input_path: str, output_path: str) -> None:
    import markdown
    md_text = _read_text_file(input_path)
    html_body = markdown.markdown(md_text, extensions=["extra", "codehilite", "tables"])
    html_doc = f"""<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<style>body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 860px; margin: 40px auto; padding: 20px; line-height: 1.7; color: #333; }}</style></head><body>{html_body}</body></html>"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_doc)


def _html_to_pdf(input_path: str, output_path: str) -> None:
    from weasyprint import HTML
    HTML(filename=input_path).write_pdf(output_path)


# ============================================================================
#  LibreOffice headless helper
# ============================================================================

def _try_libreoffice(input_path: str, output_path: str, target_fmt: str) -> bool:
    import shutil
    import subprocess
    lo = shutil.which("libreoffice") or shutil.which("soffice")
    if lo is None:
        logger.warning("未找到 LibreOffice，无法进行 %s 转换", target_fmt)
        return False
    out_dir = os.path.dirname(output_path) or "."
    try:
        subprocess.run(
            [lo, "--headless", "--convert-to", target_fmt,
             "--outdir", out_dir, input_path],
            capture_output=True, text=True, timeout=120,
        )
        lo_output = os.path.join(out_dir,
            os.path.splitext(os.path.basename(input_path))[0] + "." + target_fmt)
        if os.path.exists(lo_output) and lo_output != output_path:
            os.replace(lo_output, output_path)
        return True
    except Exception as exc:
        logger.warning("LibreOffice 转换失败: %s", exc)
        return False


# ============================================================================
#  Handler dispatch table — ALL supported conversion pairs
# ============================================================================

_HANDLERS: dict[tuple[str, str], Callable] = {
    # --- PDF → * ---
    ("pdf", "docx"): _pdf_to_docx,
    ("pdf", "txt"): _pdf_to_txt,
    ("pdf", "png"): _pdf_to_image,
    ("pdf", "jpg"): _pdf_to_image,
    ("pdf", "md"): lambda inp, out: _save_as_md(_extract_pdf_text(inp), out),
    ("pdf", "html"): lambda inp, out: _save_as_html(_extract_pdf_text(inp), out),
    ("pdf", "rtf"): lambda inp, out: _save_as_rtf(_extract_pdf_text(inp), out),
    ("pdf", "epub"): lambda inp, out: _save_text_as_epub(_extract_pdf_text(inp), out),
    ("pdf", "doc"): lambda inp, out: _chain_to_doc(inp, out, "pdf"),

    # --- DOCX → * ---
    ("docx", "pdf"): _docx_to_pdf,
    ("docx", "txt"): _docx_to_txt,
    ("docx", "html"): _docx_to_html,
    ("docx", "md"): lambda inp, out: _save_as_md(_extract_docx_text(inp), out),
    ("docx", "rtf"): lambda inp, out: _save_as_rtf(_extract_docx_text(inp), out),
    ("docx", "epub"): lambda inp, out: _save_text_as_epub(_extract_docx_text(inp), out),
    ("docx", "doc"): lambda inp, out: _chain_to_doc(inp, out, "docx"),
    # --- DOC → * (via DOCX intermediate) ---
    ("doc", "docx"): _doc_to_docx,
    ("doc", "pdf"): lambda inp, out: _chain_docx(inp, out, "pdf"),
    ("doc", "txt"): lambda inp, out: _chain_docx(inp, out, "txt"),
    ("doc", "html"): lambda inp, out: _chain_docx(inp, out, "html"),
    ("doc", "md"): lambda inp, out: _chain_docx(inp, out, "md"),
    ("doc", "rtf"): lambda inp, out: _chain_docx(inp, out, "rtf"),
    ("doc", "epub"): lambda inp, out: _chain_docx(inp, out, "epub"),
    ("doc", "doc"): lambda inp, out: shutil.copy2(inp, out),

    # --- TXT → * ---
    ("txt", "pdf"): _txt_to_pdf,
    ("txt", "docx"): _txt_to_docx,
    ("txt", "md"): lambda inp, out: _save_as_md(_read_text_file(inp), out),
    ("txt", "html"): lambda inp, out: _save_as_html(_read_text_file(inp), out),
    ("txt", "rtf"): lambda inp, out: _save_as_rtf(_read_text_file(inp), out),
    ("txt", "epub"): lambda inp, out: _save_text_as_epub(_read_text_file(inp), out),
    ("txt", "doc"): lambda inp, out: _chain_to_doc(inp, out, "txt"),

    # --- MD → * ---
    ("md", "pdf"): _md_to_pdf,
    ("md", "html"): _md_to_html,
    ("md", "txt"): lambda inp, out: _save_as_txt(_extract_md_text(inp), out),
    ("md", "docx"): lambda inp, out: _save_text_as_docx(_extract_md_text(inp), out),
    ("md", "rtf"): lambda inp, out: _save_as_rtf(_extract_md_text(inp), out),
    ("md", "epub"): lambda inp, out: _save_text_as_epub(_extract_md_text(inp), out),
    ("md", "doc"): lambda inp, out: _chain_to_doc(inp, out, "md"),

    # --- HTML → * ---
    ("html", "pdf"): _html_to_pdf,
    ("html", "txt"): lambda inp, out: _save_as_txt(_extract_html_text(inp), out),
    ("html", "docx"): lambda inp, out: _save_text_as_docx(_extract_html_text(inp), out),
    ("html", "md"): lambda inp, out: _save_as_md(_extract_html_text(inp), out),
    ("html", "rtf"): lambda inp, out: _save_as_rtf(_extract_html_text(inp), out),
    ("html", "epub"): lambda inp, out: _save_text_as_epub(_extract_html_text(inp), out),
    ("html", "doc"): lambda inp, out: _chain_to_doc(inp, out, "html"),

    # --- EPUB → * ---
    ("epub", "txt"): lambda inp, out: _save_as_txt(_extract_epub_text(inp), out),
    ("epub", "pdf"): lambda inp, out: _save_text_as_pdf(_extract_epub_text(inp), out),
    ("epub", "docx"): lambda inp, out: _save_text_as_docx(_extract_epub_text(inp), out),
    ("epub", "md"): lambda inp, out: _save_as_md(_extract_epub_text(inp), out),
    ("epub", "html"): lambda inp, out: _save_as_html(_extract_epub_text(inp), out),
    ("epub", "rtf"): lambda inp, out: _save_as_rtf(_extract_epub_text(inp), out),
    ("epub", "doc"): lambda inp, out: _chain_to_doc(inp, out, "epub"),
    ("epub", "epub"): lambda inp, out: shutil.copy2(inp, out),

    # --- RTF → * ---
    ("rtf", "txt"): lambda inp, out: _save_as_txt(_extract_rtf_text(inp), out),
    ("rtf", "pdf"): lambda inp, out: _save_text_as_pdf(_extract_rtf_text(inp), out),
    ("rtf", "docx"): lambda inp, out: _save_text_as_docx(_extract_rtf_text(inp), out),
    ("rtf", "md"): lambda inp, out: _save_as_md(_extract_rtf_text(inp), out),
    ("rtf", "html"): lambda inp, out: _save_as_html(_extract_rtf_text(inp), out),
    ("rtf", "epub"): lambda inp, out: _save_text_as_epub(_extract_rtf_text(inp), out),
    ("rtf", "rtf"): lambda inp, out: shutil.copy2(inp, out),
    ("rtf", "doc"): lambda inp, out: _chain_to_doc(inp, out, "rtf"),
}


def _raise(msg: str):
    raise RuntimeError(msg)


def _chain_to_doc(input_path: str, output_path: str, src_fmt: str) -> None:
    """Universal chain: any format → docx → doc.

    Step 1: src → docx (via _HANDLERS text extraction)
    Step 2: docx → doc (via LibreOffice headless)
    Falls back to basic text→docx→doc chain if direct route unavailable.
    """
    # Try to use existing src→docx handler
    h1 = _HANDLERS.get((src_fmt, "docx"))
    if h1 is None:
        # Fallback: extract text first
        text = _extract_any_text(src_fmt, input_path)
        h1 = lambda inp, out: _save_text_as_docx(text, out)

    tmp = tempfile.mktemp(suffix=".docx")
    try:
        h1(input_path, tmp)
        if not _try_libreoffice(tmp, output_path, "doc"):
            raise RuntimeError(f"{src_fmt.upper()} → DOC 需要安装 LibreOffice。请运行: apt-get install libreoffice-writer")
    finally:
        if os.path.exists(tmp):
            os.unlink(tmp)


def _extract_any_text(fmt: str, input_path: str) -> str:
    """Extract plain text from any supported document format."""
    extractors = {
        "pdf": _extract_pdf_text,
        "docx": _extract_docx_text,
        "doc": lambda p: "DOC requires LibreOffice for text extraction",
        "txt": _read_text_file,
        "md": _extract_md_text,
        "html": _extract_html_text,
        "rtf": _extract_rtf_text,
        "epub": _extract_epub_text,
    }
    fn = extractors.get(fmt)
    if fn:
        return fn(input_path)
    return _read_text_file(input_path)


def _chain_docx(input_path: str, output_path: str, dst_fmt: str) -> None:
    """Chain: doc → docx (temp) → dst."""
    tmp = tempfile.mktemp(suffix=".docx")
    try:
        _doc_to_docx(input_path, tmp)
        h = _HANDLERS.get(("docx", dst_fmt))
        if h:
            h(tmp, output_path)
        else:
            raise ValueError(f"Cannot chain docx→{dst_fmt}")
    finally:
        if os.path.exists(tmp):
            os.unlink(tmp)


def _chain_via_docx(src_fmt: str, dst_fmt: str, input_path: str, output_path: str) -> None:
    """Chain: src → docx → dst."""
    tmp = tempfile.mktemp(suffix=".docx")
    try:
        h1 = _HANDLERS.get((src_fmt, "docx"))
        h2 = _HANDLERS.get(("docx", dst_fmt))
        if not h1 or not h2:
            raise ValueError(f"Cannot chain {src_fmt}→docx→{dst_fmt}")
        h1(input_path, tmp)
        h2(tmp, output_path)
    finally:
        if os.path.exists(tmp):
            os.unlink(tmp)
